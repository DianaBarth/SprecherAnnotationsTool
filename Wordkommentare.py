import json
import os
import re
import time
import unicodedata
import csv
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from Eingabe import config

w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def normalisiere(text):
    text = unicodedata.normalize("NFKD", text)
    text = text.replace("\u00A0", " ")
    text = text.replace("–", "-")
    text = text.replace("—", "-")
    text = text.replace("‐", "-")
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()

def get_comments_part(part):
    for rel in part.rels.values():
        if rel.reltype == RT.COMMENTS:
            return rel._target
    package = part.package
    partname = PackURI("/word/comments.xml")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    empty_comments_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    comments_part = package.create_part(partname, content_type, empty_comments_xml)
    part.relate_to(comments_part, RT.COMMENTS)
    return comments_part

def add_comment(paragraph, text, author="Auto", initials="AU"):
    if paragraph is None:
        return

    part = paragraph.part
    comments_part = get_comments_part(part)
    if comments_part is None:
        return

    root = etree.fromstring(comments_part.blob)
    existing_ids = [int(c.get(qn("w:id"))) for c in root.findall(f'{{{w_ns}}}comment') if c.get(qn("w:id"))]
    new_id = max(existing_ids, default=0) + 1

    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), str(new_id))
    comment.set(qn('w:author'), author)
    comment.set(qn('w:initials'), initials)
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    comment.append(p)
    root.append(comment)

    comment_range_start = OxmlElement('w:commentRangeStart')
    comment_range_start.set(qn('w:id'), str(new_id))
    comment_range_end = OxmlElement('w:commentRangeEnd')
    comment_range_end.set(qn('w:id'), str(new_id))
    comment_reference = OxmlElement('w:r')
    comment_reference_run = OxmlElement('w:commentReference')
    comment_reference_run.set(qn('w:id'), str(new_id))
    comment_reference.append(comment_reference_run)

    p = paragraph._p
    p.insert(0, comment_range_start)
    p.append(comment_range_end)
    p.append(comment_reference)

    comments_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

def entferne_spezifische_kommentare(doc):
    comments_part = get_comments_part(doc.part)
    root = etree.fromstring(comments_part.blob)

    # IDs der zu entfernenden Kommentare sammeln
    zu_entfernende_ids = set()
    for comment in root.findall(f'{{{w_ns}}}comment'):
        cid = comment.get(qn('w:id'))
        texts = comment.findall(f'.//{{{w_ns}}}t')
        inhalt = ''.join([t.text or '' for t in texts]).strip()
        if re.match(r'\[[^\]]+\.\d+\]', inhalt) or not inhalt:
            if cid is not None:
                zu_entfernende_ids.add(cid)

    # Kommentare aus comments.xml entfernen
    neue_kommentare = [c for c in root.findall(f'{{{w_ns}}}comment')
                       if c.get(qn('w:id')) not in zu_entfernende_ids]
    root.clear()
    root.extend(neue_kommentare)
    comments_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # IDs auch aus dem Haupttext entfernen
    for para in doc.paragraphs:
        p = para._p
        remove = []
        for el in list(p):
            cid = el.get(qn('w:id'))
            if el.tag in {qn('w:commentRangeStart'), qn('w:commentRangeEnd')} and cid in zu_entfernende_ids:
                remove.append(el)
            elif el.tag == qn('w:r'):
                for child in el:
                    if child.tag == qn('w:commentReference') and child.get(qn('w:id')) in zu_entfernende_ids:
                        remove.append(el)
        for el in remove:
            p.remove(el)

    print(f"✅ {len(zu_entfernende_ids)} Kommentare vollständig entfernt.")

def speichere_csv_sicher(pfad, daten):
    max_versuche = 3
    for versuch in range(max_versuche):
        if ist_datei_gesperrt(pfad):
            print(f"⛔ CSV-Datei '{pfad}' ist geöffnet. Bitte schließen und mit [Enter] bestätigen.")
            input()
            time.sleep(1)
        else:
            try:
                with open(pfad, mode='w', encoding='utf-8', newline='') as csvfile:
                    writer = csv.writer(csvfile, delimiter=';')
                    writer.writerow(["Kapitelname", "UnterkapitelNr", "Wortanzahl"])
                    writer.writerows(daten)
                print(f"✅ CSV-Datei gespeichert unter: {pfad}")
                break
            except Exception as e:
                print(f"❌ Fehler beim Speichern der CSV-Datei: {e}")
    else:
        print("❌ CSV-Datei konnte nicht gespeichert werden. Bitte manuell schließen.")


def verarbeite_kapitel(doc, kapitel_config, csv_pfad=None):
    kapitel_liste = kapitel_config["kapitel_liste"]
    kapitel_trenner = kapitel_config["Kapitel_trenner"]

    aktuelles_kapitel_text = None
    unterkapitel_nr = 1
    wort_zaehler = 0
    start_absatz = None
    absatz_iter = iter(doc.paragraphs)
    csv_daten = []

    for absatz in absatz_iter:
        text = absatz.text.strip()
        if text in kapitel_liste:
            if aktuelles_kapitel_text is not None and start_absatz is not None:
                kommentar_text = f"[{aktuelles_kapitel_text}.{unterkapitel_nr}] mit {wort_zaehler} Wörtern"
                add_comment(start_absatz, kommentar_text)
                csv_daten.append([aktuelles_kapitel_text, unterkapitel_nr, wort_zaehler])
            aktuelles_kapitel_text = text
            unterkapitel_nr = 1
            wort_zaehler = 0
            start_absatz = absatz
            continue

        if text == kapitel_trenner and aktuelles_kapitel_text is not None:
            kommentar_text = f"[{aktuelles_kapitel_text}.{unterkapitel_nr}] mit {wort_zaehler} Wörtern"
            add_comment(start_absatz, kommentar_text)
            csv_daten.append([aktuelles_kapitel_text, unterkapitel_nr, wort_zaehler])
            unterkapitel_nr += 1
            wort_zaehler = 0
            start_absatz = absatz
            continue

        if aktuelles_kapitel_text is not None:
            wort_zaehler += len(text.split())

    if aktuelles_kapitel_text is not None and start_absatz is not None:
        kommentar_text = f"[{aktuelles_kapitel_text}.{unterkapitel_nr}] mit {wort_zaehler} Wörtern (final)"
        add_comment(start_absatz, kommentar_text)
        csv_daten.append([aktuelles_kapitel_text, unterkapitel_nr, wort_zaehler])

    if csv_pfad:
        speichere_csv_sicher(csv_pfad, csv_daten)

def ist_datei_gesperrt(pfad):
    try:
        with open(pfad, 'a'):
            return False
    except PermissionError:
        return True

# --- Hauptausführung ---
GLOBALORDNER = config.GLOBALORDNER
eingabe_ordner = GLOBALORDNER['Eingabe']
parent_ordner = os.path.dirname(eingabe_ordner)
kapitel_config_path = os.path.join(parent_ordner, "kapitel_config.json")

projekt_root = os.path.abspath(os.path.join(eingabe_ordner, "..", ".."))
projektname = os.path.basename(parent_ordner)
docx_path = os.path.join(os.path.dirname(projekt_root), f"{projektname}.docx")

zielpfad = docx_path.replace(".docx", "_kommentiert.docx")
csv_path = docx_path.replace(".docx", "_wortzaehler.csv")

print(f"📄 Öffne Dokument: {docx_path}")
doc = Document(docx_path)

with open(kapitel_config_path, "r", encoding="utf-8") as f:
    kapitel_config = json.load(f)

entferne_spezifische_kommentare(doc)
#verarbeite_kapitel(doc, kapitel_config, csv_path)

max_versuche = 3
for versuch in range(max_versuche):
    if ist_datei_gesperrt(zielpfad):
        print(f"⛔ Datei '{zielpfad}' ist geöffnet. Bitte schließen und mit [Enter] bestätigen.")
        input()
        time.sleep(1)
    else:
        try:
            doc.save(zielpfad)
            print(f"✅ Fertig: Datei mit Kommentaren gespeichert:\n{zielpfad}")
        except Exception as e:
            print(f"❌ Fehler beim Speichern: {e}")
        break
else:
    print("❌ Datei konnte nach mehreren Versuchen nicht gespeichert werden. Bitte manuell schließen.")
