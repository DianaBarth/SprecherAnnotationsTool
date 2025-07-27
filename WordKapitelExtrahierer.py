import json
import os
from docx import Document
from Eingabe import config

# === Projektstruktur ===
GLOBALORDNER = config.GLOBALORDNER
eingabe_ordner = GLOBALORDNER['Eingabe']

parent_ordner = os.path.dirname(eingabe_ordner)
kapitel_config_path = os.path.join(parent_ordner, "kapitel_config.json")

projekt_root = os.path.abspath(os.path.join(eingabe_ordner, "..", ".."))
projektname = os.path.basename(parent_ordner)
docx_path = os.path.join(os.path.dirname(projekt_root), f"{projektname}.docx")

ausgabe_ordner = os.path.join(parent_ordner, "Kapitel")
os.makedirs(ausgabe_ordner, exist_ok=True)

# === Kapitelüberschriften laden ===
with open(kapitel_config_path, encoding='utf-8') as f:
    kapitel_liste = json.load(f)["kapitel_liste"]

# === Original-Dokument laden ===
original_doc = Document(docx_path)
absätze = original_doc.paragraphs

# === Hilfsfunktion: Kapitelstart-Indexe finden ===
def finde_kapitel_start_indexe(absätze, kapitel_liste):
    indexe = []
    for kapitel in kapitel_liste:
        for i, p in enumerate(absätze):
            if p.text.strip() == kapitel.strip():
                indexe.append(i)
                break
        else:
            print(f"⚠️ Kapitel nicht gefunden: {kapitel}")
            indexe.append(None)
    return indexe

kapitel_indexe = finde_kapitel_start_indexe(absätze, kapitel_liste)

# === Kapitel extrahieren ===
for i in range(len(kapitel_liste)):
    start_idx = kapitel_indexe[i]
    if start_idx is None:
        continue
    end_idx = kapitel_indexe[i + 1] if i + 1 < len(kapitel_liste) and kapitel_indexe[i + 1] is not None else len(absätze)

    # Dokument neu laden (Original bleibt unberührt)
    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    # Absätze NACH dem Kapitel löschen
    for j in reversed(range(end_idx, len(paragraphs))):
        p = paragraphs[j]._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Absätze VOR dem Kapitel löschen
    for j in reversed(range(0, start_idx)):
        p = paragraphs[j]._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Datei speichern
    safe_name = f"{i:02d}_{kapitel_liste[i].replace(' ', '_').replace('/', '-')}.docx"
    save_path = os.path.join(ausgabe_ordner, safe_name)
    doc.save(save_path)
    print(f"✅ Kapitel gespeichert: {save_path}")
